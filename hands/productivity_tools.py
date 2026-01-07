"""
Productivity Tools - Email, Documents, Calendar, Reminders
Everyday productivity automation for GLOW
"""

import os
import json
import subprocess
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional, List, Dict
import pyautogui
import time


# ===== EMAIL TOOLS =====

def draft_email(to: str, subject: str, body: str, cc: Optional[str] = None) -> str:
    """
    Draft an email in default mail client (Outlook/Gmail)

    Args:
        to: Recipient email address
        subject: Email subject
        body: Email body text
        cc: CC recipients (optional)

    Returns:
        Success message
    """
    try:
        # Create mailto URL
        mailto_url = f"mailto:{to}?subject={subject}&body={body}"
        if cc:
            mailto_url += f"&cc={cc}"

        # Open in default mail client
        os.startfile(mailto_url)
        time.sleep(2)

        return f"Email draft opened: To={to}, Subject='{subject}'"
    except Exception as e:
        return f"Error drafting email: {str(e)}"


def check_screen_for_text(search_text: str) -> str:
    """
    Take screenshot and check if specific text appears on screen
    Useful for checking emails, notifications, etc.

    Args:
        search_text: Text to search for on screen

    Returns:
        Found/Not found message
    """
    try:
        import pytesseract
        from PIL import Image

        # Take screenshot
        screenshot = pyautogui.screenshot()

        # OCR to extract text
        text = pytesseract.image_to_string(screenshot)

        if search_text.lower() in text.lower():
            return f"Found '{search_text}' on screen"
        else:
            return f"'{search_text}' not found on screen"
    except Exception as e:
        return f"Error checking screen: {str(e)}"


# ===== DOCUMENT TOOLS =====

def open_word() -> str:
    """Open Microsoft Word"""
    try:
        subprocess.Popen(["start", "winword"], shell=True)
        time.sleep(3)
        return "Microsoft Word opened"
    except Exception as e:
        return f"Error opening Word: {str(e)}"


def open_excel() -> str:
    """Open Microsoft Excel"""
    try:
        subprocess.Popen(["start", "excel"], shell=True)
        time.sleep(3)
        return "Microsoft Excel opened"
    except Exception as e:
        return f"Error opening Excel: {str(e)}"


def open_powerpoint() -> str:
    """Open Microsoft PowerPoint"""
    try:
        subprocess.Popen(["start", "powerpnt"], shell=True)
        time.sleep(3)
        return "Microsoft PowerPoint opened"
    except Exception as e:
        return f"Error opening PowerPoint: {str(e)}"


def create_word_document(filename: str, content: str, save_path: Optional[str] = None) -> str:
    """
    Create a Word document with content using python-docx

    Args:
        filename: Document filename (without .docx)
        content: Text content to write (can include formatting markers)
        save_path: Where to save (default: Desktop)

    Returns:
        Success message with file path
    """
    try:
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            return "Error: python-docx not installed. Run: pip install python-docx"

        # Create document
        doc = Document()

        # Parse content for structured document
        # Check if content has section markers
        if "TITLE:" in content or "SECTION:" in content:
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith("TITLE:"):
                    # Add title
                    title = line.replace("TITLE:", "").strip()
                    heading = doc.add_heading(title, level=0)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                elif line.startswith("SECTION:"):
                    # Add section heading
                    section_name = line.replace("SECTION:", "").strip()
                    doc.add_heading(section_name, level=1)

                elif line.startswith("- ") or line.startswith("• "):
                    # Add bullet point
                    bullet_text = line[2:].strip()
                    doc.add_paragraph(bullet_text, style='List Bullet')

                else:
                    # Add regular paragraph
                    doc.add_paragraph(line)
        else:
            # Simple content - just add as paragraphs
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)

        # Determine save path
        if not filename.endswith('.docx'):
            filename += '.docx'

        if save_path:
            full_path = os.path.join(save_path, filename)
        else:
            # Default to Desktop
            desktop = os.path.join(os.path.expanduser("~"), "Desktop")
            full_path = os.path.join(desktop, filename)

        # Save document
        doc.save(full_path)

        return f"Word document created: {full_path}"
    except Exception as e:
        return f"Error creating Word document: {str(e)}"


def create_excel_spreadsheet(filename: str, data: List[List[str]], save_path: Optional[str] = None) -> str:
    """
    Create an Excel spreadsheet with data

    Args:
        filename: Spreadsheet filename (without .xlsx)
        data: 2D list of cell values [[row1], [row2], ...]
        save_path: Where to save (default: Documents folder)

    Returns:
        Success message with file path
    """
    try:
        # Open Excel
        open_excel()
        time.sleep(2)

        # Enter data
        for row in data:
            for cell in row:
                pyautogui.write(str(cell), interval=0.01)
                pyautogui.press('tab')
            pyautogui.press('enter')

        time.sleep(0.5)

        # Save file (Ctrl+S)
        pyautogui.hotkey('ctrl', 's')
        time.sleep(1)

        # Type filename
        if not filename.endswith('.xlsx'):
            filename += '.xlsx'

        if save_path:
            full_path = os.path.join(save_path, filename)
        else:
            docs_path = os.path.expanduser("~/Documents")
            full_path = os.path.join(docs_path, filename)

        pyautogui.write(full_path, interval=0.01)
        pyautogui.press('enter')

        return f"Excel spreadsheet created: {full_path}"
    except Exception as e:
        return f"Error creating Excel spreadsheet: {str(e)}"


# ===== NOTEPAD & TEXT EDITING =====

def open_notepad(content: Optional[str] = None) -> str:
    """
    Open Notepad and optionally type content

    Args:
        content: Text to type in Notepad (optional)

    Returns:
        Success message
    """
    try:
        subprocess.Popen(["notepad.exe"])
        time.sleep(1)

        if content:
            time.sleep(0.5)
            pyautogui.write(content, interval=0.01)

        return f"Notepad opened{' with content' if content else ''}"
    except Exception as e:
        return f"Error opening Notepad: {str(e)}"


def save_notepad(filename: str, save_path: Optional[str] = None) -> str:
    """
    Save current Notepad file

    Args:
        filename: Filename to save as
        save_path: Where to save (default: Documents)

    Returns:
        Success message
    """
    try:
        pyautogui.hotkey('ctrl', 's')
        time.sleep(0.5)

        if save_path:
            full_path = os.path.join(save_path, filename)
        else:
            docs_path = os.path.expanduser("~/Documents")
            full_path = os.path.join(docs_path, filename)

        pyautogui.write(full_path, interval=0.01)
        pyautogui.press('enter')

        return f"Notepad saved: {full_path}"
    except Exception as e:
        return f"Error saving Notepad: {str(e)}"


# ===== CALENDAR & REMINDERS =====

def create_reminder(title: str, time_minutes: int, message: Optional[str] = None) -> str:
    """
    Create a Windows notification reminder

    Args:
        title: Reminder title
        time_minutes: Minutes from now to remind
        message: Optional reminder message

    Returns:
        Success message
    """
    try:
        from pathlib import Path
        reminder_file = Path.home() / ".glow_reminders.json"

        # Load existing reminders
        reminders = []
        if reminder_file.exists():
            with open(reminder_file, 'r') as f:
                reminders = json.load(f)

        # Add new reminder
        remind_time = datetime.now() + timedelta(minutes=time_minutes)
        reminders.append({
            "title": title,
            "message": message or "",
            "time": remind_time.isoformat(),
            "completed": False
        })

        # Save reminders
        with open(reminder_file, 'w') as f:
            json.dump(reminders, f, indent=2)

        return f"Reminder set: '{title}' in {time_minutes} minutes ({remind_time.strftime('%I:%M %p')})"
    except Exception as e:
        return f"Error creating reminder: {str(e)}"


def list_reminders() -> str:
    """List all active reminders"""
    try:
        from pathlib import Path
        reminder_file = Path.home() / ".glow_reminders.json"

        if not reminder_file.exists():
            return "No reminders found"

        with open(reminder_file, 'r') as f:
            reminders = json.load(f)

        active = [r for r in reminders if not r.get('completed')]

        if not active:
            return "No active reminders"

        result = "Active Reminders:\n"
        for i, r in enumerate(active, 1):
            remind_time = datetime.fromisoformat(r['time'])
            result += f"{i}. {r['title']} - {remind_time.strftime('%I:%M %p on %b %d')}\n"
            if r.get('message'):
                result += f"   Message: {r['message']}\n"

        return result.strip()
    except Exception as e:
        return f"Error listing reminders: {str(e)}"


# ===== SEARCH & WEB RESEARCH =====

def search_web(query: str, engine: str = "google") -> str:
    """
    Search the web using specified search engine

    Args:
        query: Search query
        engine: Search engine (google/bing/duckduckgo)

    Returns:
        Success message
    """
    try:
        engines = {
            "google": f"https://www.google.com/search?q={query}",
            "bing": f"https://www.bing.com/search?q={query}",
            "duckduckgo": f"https://duckduckgo.com/?q={query}"
        }

        url = engines.get(engine.lower(), engines["google"])
        os.startfile(url)

        return f"Searching {engine} for: {query}"
    except Exception as e:
        return f"Error searching web: {str(e)}"


def open_website(url: str) -> str:
    """
    Open a website in default browser

    Args:
        url: Website URL

    Returns:
        Success message
    """
    try:
        if not url.startswith('http'):
            url = 'https://' + url

        os.startfile(url)
        time.sleep(2)

        return f"Opened website: {url}"
    except Exception as e:
        return f"Error opening website: {str(e)}"


# ===== FILE ORGANIZATION =====

def organize_downloads() -> str:
    """Organize Downloads folder by file type"""
    try:
        downloads = Path.home() / "Downloads"

        # Create organization folders
        folders = {
            "Documents": ['.pdf', '.doc', '.docx', '.txt', '.xlsx', '.pptx'],
            "Images": ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg'],
            "Videos": ['.mp4', '.avi', '.mkv', '.mov', '.wmv'],
            "Archives": ['.zip', '.rar', '.7z', '.tar', '.gz'],
            "Executables": ['.exe', '.msi'],
            "Code": ['.py', '.js', '.html', '.css', '.java', '.cpp']
        }

        moved_count = 0
        for folder_name, extensions in folders.items():
            folder_path = downloads / folder_name
            folder_path.mkdir(exist_ok=True)

            for file in downloads.iterdir():
                if file.is_file() and file.suffix.lower() in extensions:
                    dest = folder_path / file.name
                    file.rename(dest)
                    moved_count += 1

        return f"Organized {moved_count} files in Downloads folder"
    except Exception as e:
        return f"Error organizing downloads: {str(e)}"


def find_large_files(min_size_mb: int = 100, search_path: Optional[str] = None) -> str:
    """
    Find large files on system

    Args:
        min_size_mb: Minimum file size in MB
        search_path: Path to search (default: user home)

    Returns:
        List of large files
    """
    try:
        search_dir = Path(search_path) if search_path else Path.home()
        min_bytes = min_size_mb * 1024 * 1024

        large_files = []
        for file in search_dir.rglob('*'):
            if file.is_file():
                try:
                    if file.stat().st_size >= min_bytes:
                        size_mb = file.stat().st_size / (1024 * 1024)
                        large_files.append((str(file), size_mb))
                except (PermissionError, OSError):
                    continue

        if not large_files:
            return f"No files larger than {min_size_mb}MB found"

        # Sort by size
        large_files.sort(key=lambda x: x[1], reverse=True)

        result = f"Found {len(large_files)} files larger than {min_size_mb}MB:\n"
        for path, size in large_files[:10]:  # Top 10
            result += f"  {size:.1f}MB - {path}\n"

        return result.strip()
    except Exception as e:
        return f"Error finding large files: {str(e)}"


# ===== QUICK ACTIONS =====

def open_calculator() -> str:
    """Open Windows Calculator"""
    try:
        subprocess.Popen(["calc.exe"])
        return "Calculator opened"
    except Exception as e:
        return f"Error opening calculator: {str(e)}"


def open_calendar() -> str:
    """Open Windows Calendar"""
    try:
        subprocess.Popen(["start", "outlookcal:"], shell=True)
        return "Calendar opened"
    except Exception as e:
        return f"Error opening calendar: {str(e)}"


def open_task_manager() -> str:
    """Open Windows Task Manager"""
    try:
        subprocess.Popen(["taskmgr.exe"])
        return "Task Manager opened"
    except Exception as e:
        return f"Error opening Task Manager: {str(e)}"


def open_file_explorer(path: Optional[str] = None) -> str:
    """
    Open File Explorer at specified path

    Args:
        path: Folder path to open (default: This PC)

    Returns:
        Success message
    """
    try:
        if path:
            subprocess.Popen(f'explorer "{path}"')
        else:
            subprocess.Popen('explorer')

        return f"File Explorer opened{f' at {path}' if path else ''}"
    except Exception as e:
        return f"Error opening File Explorer: {str(e)}"


def empty_recycle_bin() -> str:
    """Empty Windows Recycle Bin"""
    try:
        import winshell
        winshell.recycle_bin().empty(confirm=False, show_progress=False, sound=False)
        return "Recycle Bin emptied"
    except ImportError:
        return "Error: winshell package not installed. Run: pip install winshell"
    except Exception as e:
        return f"Error emptying Recycle Bin: {str(e)}"


# ===== LIVE TYPING TOOLS FOR WORD/EXCEL =====

def open_word_and_type(content: str, save_as: Optional[str] = None) -> str:
    """
    Open Microsoft Word and type content LIVE (visible to user)

    Args:
        content: Text to type into Word
        save_as: Optional filename to save (on desktop)

    Returns:
        Status message
    """
    try:
        # Open Word
        subprocess.Popen(['start', 'winword'], shell=True)
        time.sleep(3)  # Wait for Word to open

        # Type content live (user will see it being typed)
        pyautogui.write(content, interval=0.03)  # 30ms between chars

        if save_as:
            # Save file (Ctrl+S)
            time.sleep(0.5)
            pyautogui.hotkey('ctrl', 's')
            time.sleep(1)

            # Type filename
            pyautogui.write(save_as, interval=0.05)
            pyautogui.press('enter')
            time.sleep(0.5)

            return f"Typed into Word and saved as: {save_as}"

        return f"Typed into Word: {len(content)} characters"

    except Exception as e:
        return f"Error: {str(e)}"


def open_excel_and_enter_data(data: List[List[str]], save_as: Optional[str] = None) -> str:
    """
    Open Excel and enter data LIVE into cells

    Args:
        data: 2D list of data [[row1], [row2], ...]
        save_as: Optional filename to save

    Returns:
        Status message
    """
    try:
        # Open Excel
        subprocess.Popen(['start', 'excel'], shell=True)
        time.sleep(3)  # Wait for Excel to open

        # Click to ensure we're in cell A1
        time.sleep(0.5)

        # Enter data live
        for row_idx, row in enumerate(data):
            for col_idx, cell in enumerate(row):
                # Type cell content
                pyautogui.write(str(cell), interval=0.03)

                # Move to next cell (Tab for next column)
                if col_idx < len(row) - 1:
                    pyautogui.press('tab')
                    time.sleep(0.1)

            # Move to next row (Enter)
            if row_idx < len(data) - 1:
                pyautogui.press('enter')
                # Move back to column A
                for _ in range(len(row) - 1):
                    pyautogui.press('left')
                time.sleep(0.1)

        if save_as:
            time.sleep(0.5)
            pyautogui.hotkey('ctrl', 's')
            time.sleep(1)
            pyautogui.write(save_as, interval=0.05)
            pyautogui.press('enter')
            return f"Entered {len(data)} rows into Excel and saved as: {save_as}"

        return f"Entered {len(data)} rows into Excel"

    except Exception as e:
        return f"Error: {str(e)}"


def type_in_active_window(text: str, interval: float = 0.03) -> str:
    """
    Type text in the currently active window (live typing)

    Args:
        text: Text to type
        interval: Delay between keystrokes (seconds)

    Returns:
        Status message
    """
    try:
        time.sleep(0.5)  # Small delay to ensure window is ready
        pyautogui.write(text, interval=interval)
        return f"Typed {len(text)} characters into active window"
    except Exception as e:
        return f"Error: {str(e)}"
